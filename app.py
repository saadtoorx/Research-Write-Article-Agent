import streamlit as st
import os
import warnings
from crewai import Crew
from utils import get_openai_api_key
import io
from docx import Document

from agents import create_agents
from tasks import create_tasks

# Suppress warnings
warnings.filterwarnings('ignore')

# Page configuration
st.set_page_config(
    page_title="Research & Write Article Agent",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Playwrite+AU+QLD:wght@100..400&display=swap');
    @import url('https://fonts.googleapis.com/css2?family=Poppins:ital,wght@0,100;0,200;0,300;0,400;0,500;0,600;0,700;0,800;0,900;1,100;1,200;1,300;1,400;1,500;1,600;1,700;1,800;1,900&display=swap');
     
    /* Apply custom fonts */
    .stApp {
        font-family: 'Inter', sans-serif;
    }

    .main-header {
        font-family: "Playwrite AU QLD", cursive !important;
        font-optical-sizing: auto !important;
        font-weight: <weight> !important;
        font-style: normal;
        text-align: center;
        margin-bottom: 2rem;
    }
    .sub-header {
        font-family: "Poppins", sans-serif !important;
        font-size: 1.2rem;
        color: #666;
        text-align: center;
    }
    .result-box {
        background-color: #f8f9fa;
        border: 1px solid #dee2e6;
        border-radius: 10px;
        padding: 20px;
        margin: 10px 0;
    }
    .agent-info {
        background-color: #e3f2fd;
        border-left: 4px solid #2196f3;
        padding: 15px;
        margin: 10px 0;
        border-radius: 5px;
    }

    .output_blog {
        font-family: "Poppins", sans-serif !important;
        font-size: 1.1rem;
        color: #1e1e1e;
        background-color: #fdfdfd !important;
        padding: 2rem 4rem;
        border-radius: 10px;
    }
</style>
""", unsafe_allow_html=True)

def setup_openai_api():
    """Setup OpenAI API key"""
    try:
        api_key = get_openai_api_key()
        os.environ["OPENAI_API_KEY"] = api_key
        return True
    except Exception as e:
        st.error(f"Error setting up OpenAI API: {str(e)}")
        return False



def generate_article(topic):
    """Generate article using CrewAI"""
    try:
        # Create agents and tasks
        planner, writer, editor = create_agents()
        plan, write, edit = create_tasks(planner, writer, editor)
        
        # Create crew
        crew = Crew(
            agents=[planner, writer, editor],
            tasks=[plan, write, edit],
            verbose=True
        )
        
        # Run the crew
        result = crew.kickoff(inputs={"topic": topic})
        return result
        
    except Exception as e:
        st.error(f"Error generating article: {str(e)}")
        return None

def main():
    # Header
    st.markdown('<h1 class="main-header">📝 Research & Write Article Agent</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Planner - Writer - Editor</p>', unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Key setup
        st.subheader("OpenAI API Key")
        api_key_input = st.text_input(
            "Enter your OpenAI API Key",
            type="password",
            help="Your OpenAI API key is required to run this application"
        )
        
        if api_key_input:
            os.environ["OPENAI_API_KEY"] = api_key_input
            st.success("✅ API Key configured")
        
        # Agent information
        st.subheader("🤖 AI Agents")
        st.markdown("""
        **Content Planner**: Researches and plans content structure\n
        **Content Writer**: Creates engaging articles\n
        **Editor**: Reviews and polishes content
        """)
        
        # Features
        st.subheader("✨ Features")
        st.markdown("""
        • Comprehensive research\n
        • SEO optimization\n
        • Grammar checking\n
        • Style consistency\n
        • Markdown formatting
        """)
    
    # Main content area
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        # Topic input
        st.subheader("🎯 Enter Your Topic")
        topic = st.text_input(
            "What would you like to write about?",
            placeholder="e.g., AI in Healthcare, Climate Change Solutions, Digital Marketing Trends...",
            help="Enter a specific topic or subject for your article"
        )
        
        # Generate button
        if st.button("🚀 Generate Article", type="primary", use_container_width=True):
            if not topic:
                st.warning("Please enter a topic to generate an article.")
            elif not os.getenv("OPENAI_API_KEY"):
                st.error("Please configure your OpenAI API key in the sidebar.")
            else:
                with st.spinner("🤖 AI agents are working on your article..."):
                    # Progress bar
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    # Simulate progress
                    for i in range(3):
                        if i == 0:
                            status_text.text("📋 Planning content structure...")
                        elif i == 1:
                            status_text.text("✍️ Writing the article...")
                        else:
                            status_text.text("🔍 Editing and polishing...")
                        progress_bar.progress((i + 1) * 33)
                    
                    # Generate article
                    result = generate_article(topic)
                    
                    if result:
                        progress_bar.progress(100)
                        status_text.text("✅ Article generated successfully!")

                        # --- Display Success Message ---
                        st.success("🎉 Your article is ready!")

                        # --- BLOG-STYLE OUTPUT (Minimalist & Clean) ---
                        st.markdown("## 📰 Blog Article", unsafe_allow_html=True)

                        # Convert Markdown-style headings (#, ##) into bold text
                        def clean_markdown(text):
                            lines = text.split("\n")
                            formatted_lines = []
                            for line in lines:
                                line = line.strip()
                                if line.startswith("### "):
                                    line = f"<strong style='font-size:1.2em'>{line[4:].strip()}</strong>"
                                elif line.startswith("## "):
                                    line = f"<strong style='font-size:1.4em'>{line[3:].strip()}</strong>"
                                elif line.startswith("# "):
                                    line = f"<strong style='font-size:1.6em'>{line[2:].strip()}</strong>"
                                formatted_lines.append(line)
                            return "<br><br>".join(formatted_lines)

                        styled_article = clean_markdown(result.raw)

                        # --- Display Blog-Styled Article ---
                        st.markdown(
                            f"""
                            <div class="output_blog">
                                {styled_article}
                            </div>
                            """,
                            unsafe_allow_html=True
                        )

                        # --- DOWNLOAD BUTTON ---
                        st.markdown("<div style='margin-top: 2em;'></div>", unsafe_allow_html=True)
                        col1, col2= st.columns([1, 1])

                        with col1:
                            try:
                                st.download_button(
                                    label="📥 Download Article (Markdown)",
                                    data=result.raw,
                                    file_name=f"{topic.replace(' ', '_').lower()}_article.md",
                                    mime="text/markdown"
                                )
                            except Exception as e:
                                st.warning(f"Markdown download unavailable: {e}")
                        
                        with col2:
                            # --- DOCX DOWNLOAD FUNCTION & BUTTON ---
                            def create_docx(content: str, topic: str):
                                doc = Document()
                                doc.add_heading(topic, 0)

                                for line in content.split("\n"):
                                    if line.strip():
                                        doc.add_paragraph(line.strip())
                                
                                buffer = io.BytesIO()
                                doc.save(buffer)
                                buffer.seek(0)
                                return buffer

                            docx_file = create_docx(result.raw, topic)
                            st.download_button(
                                label="📄 Download as Word (.docx)",
                                data=docx_file,
                                file_name=f"{topic.replace(' ', '_').lower()}_article.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                        # --- ARTICLE METADATA ---
                        with st.expander("📊 Article Details"):
                            st.markdown(f"""
                            - **📝 Topic:** `{topic}`
                            - **🤖 Generated by:** CrewAI Multi-Agent System  
                            - **📄 Format:** Blog-style (Markdown converted to styled HTML)  
                            - **🔢 Length:** `{len(result.raw):,}` characters
                            """)

                    else:
                        st.error("❌ Failed to generate article. Please check your API key and try again.")
    
    # Footer
    st.markdown("---")
    st.markdown(
        "<p style='text-align: center; color: #666;'>Built with CrewAI and Streamlit | Research & Write Article Agent</p>"
        "<p style='text-align: center; color: #666;'>Saad Toor | @saadtoorx</p>",
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main() 